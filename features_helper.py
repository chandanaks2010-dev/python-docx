import streamlit as st
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from io import BytesIO

# ── Shared constants ──────────────────────────────────────────────────────────

ALIGN_MAP = {
    "Left":    WD_PARAGRAPH_ALIGNMENT.LEFT,
    "Center":  WD_PARAGRAPH_ALIGNMENT.CENTER,
    "Right":   WD_PARAGRAPH_ALIGNMENT.RIGHT,
    "Justify": WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
}

HEADING_SIZE = {1: "2em", 2: "1.5em", 3: "1.17em"}

# ── Private helpers ───────────────────────────────────────────────────────────

def _apply_run_fmt(run, block):
    """Apply bold/italic/underline/size/color from a block to a docx run."""
    run.bold      = block.get("bold", False)
    run.italic    = block.get("italic", False)
    run.underline = block.get("underline", False)
    run.font.size = Pt(block.get("font_size", 12))
    try:
        c = block.get("font_color", "#000000")
        run.font.color.rgb = RGBColor(int(c[1:3], 16), int(c[3:5], 16), int(c[5:7], 16))
    except Exception:
        pass


def _css_style(block):
    """Build an inline CSS style string from block formatting for preview."""
    align = block.get("alignment", "Left")
    css_align = "justify" if align == "Justify" else align.lower()
    style = f"text-align: {css_align}; "
    if block.get("bold"):      style += "font-weight: bold; "
    if block.get("italic"):    style += "font-style: italic; "
    if block.get("underline"): style += "text-decoration: underline; "
    style += f"font-size: {block.get('font_size', 12)}px; "
    style += f"color: {block.get('font_color', '#000000')};"
    return style


def init_session_state():
    if "doc" not in st.session_state:
        st.session_state.doc = Document()
        st.session_state.content_blocks = []
        st.session_state.bold = False
        st.session_state.italic = False
        st.session_state.underline = False
        st.session_state.font_size = 12
        st.session_state.font_color = "#000000"


def save_document():
    """Feature 10: Save/Download Document"""
    rebuild_document()
    buffer = BytesIO()
    st.session_state.doc.save(buffer)
    buffer.seek(0)
    return buffer


def new_document():
    """Create a fresh Document and clear blocks"""
    st.session_state.doc = Document()
    st.session_state.content_blocks = []


def add_block(block_type, content, **kwargs):
    """Add content block"""
    st.session_state.content_blocks.append({
        "type": block_type,
        "content": content,
        **kwargs
    })


def rebuild_document():
    """Rebuild document from blocks"""
    st.session_state.doc = Document()
    for block in st.session_state.content_blocks:
        if block["type"] == "heading":
            p = st.session_state.doc.add_heading(block["content"], level=block.get("level", 1))
            p.alignment = ALIGN_MAP.get(block.get("alignment", "Left"), WD_PARAGRAPH_ALIGNMENT.LEFT)
            for run in p.runs:
                _apply_run_fmt(run, block)
        elif block["type"] == "paragraph":
            p = st.session_state.doc.add_paragraph()
            run = p.add_run(block.get("content", ""))
            _apply_run_fmt(run, block)
            p.alignment = ALIGN_MAP.get(block.get("alignment", "Left"), WD_PARAGRAPH_ALIGNMENT.LEFT)
        elif block["type"] in ("bullet_list", "numbered_list"):
            list_style = "List Bullet" if block["type"] == "bullet_list" else "List Number"
            for item in block.get("content", []):
                p = st.session_state.doc.add_paragraph(item, style=list_style)
                p.alignment = ALIGN_MAP.get(block.get("alignment", "Left"), WD_PARAGRAPH_ALIGNMENT.LEFT)
                for run in p.runs:
                    _apply_run_fmt(run, block)
        elif block["type"] == "table":
            rows = block.get("rows", max(1, len(block.get("data", []))))
            cols = block.get("cols", max(1, max((len(r) for r in block.get("data", [])), default=1)))
            table = st.session_state.doc.add_table(rows=rows, cols=cols)
            for ri, row_data in enumerate(block.get("data", [])):
                for cj, cell_content in enumerate(row_data):
                    cell = table.cell(ri, cj)
                    cell.text = ""
                    p = cell.paragraphs[0]
                    run = p.add_run(str(cell_content))
                    _apply_run_fmt(run, block)
                    p.alignment = ALIGN_MAP.get(block.get("alignment", "Left"), WD_PARAGRAPH_ALIGNMENT.LEFT)
        elif block["type"] == "image":
            try:
                width_in = block.get("width_in", 3)
                if width_in:
                    st.session_state.doc.add_picture(block["content"], width=Inches(width_in))
                else:
                    st.session_state.doc.add_picture(block["content"])
                st.session_state.doc.paragraphs[-1].alignment = ALIGN_MAP.get(block.get("alignment", "Left"), WD_PARAGRAPH_ALIGNMENT.LEFT)
            except Exception:
                pass
        elif block["type"] == "page_break":
            st.session_state.doc.add_page_break()


def render_preview():
    """Show live preview"""
    with st.container():
        if not st.session_state.content_blocks:
            st.info("👈 Add content from features above")
            return

        for i, block in enumerate(st.session_state.content_blocks):
            with st.container():
                if block["type"] == "heading":
                    level = block.get("level", 1)
                    style = _css_style(block)
                    style += f"font-size: {HEADING_SIZE.get(level, '2em')}; font-weight: bold; "
                    st.markdown(f"<p style='{style} margin: 0 0 8px 0;'>{block['content']}</p>", unsafe_allow_html=True)
                elif block["type"] == "paragraph":
                    st.markdown(f"<p style='{_css_style(block)}'>{block['content']}</p>", unsafe_allow_html=True)
                elif block["type"] == "bullet_list":
                    style = _css_style(block)
                    for item in block.get("content", []):
                        st.markdown(f"<p style='margin:2px 0; {style}'>• {item}</p>", unsafe_allow_html=True)
                elif block["type"] == "numbered_list":
                    style = _css_style(block)
                    for idx, item in enumerate(block.get("content", []), 1):
                        st.markdown(f"<p style='margin:2px 0; {style}'>{idx}. {item}</p>", unsafe_allow_html=True)
                elif block["type"] == "table":
                    data = block.get("data")
                    if data:
                        cell_style = _css_style(block)
                        html = "<table style='border-collapse: collapse; width: 100%;'>"
                        for row in data:
                            html += "<tr>" + "".join(
                                f"<td style='border:1px solid #ddd; padding:6px; {cell_style}'>{str(cell)}</td>"
                                for cell in row
                            ) + "</tr>"
                        html += "</table>"
                        st.markdown(html, unsafe_allow_html=True)
                    else:
                        st.write("📊 Table inserted (no data)")
                elif block["type"] == "image":
                    content = block.get("content")
                    if content is not None:
                        try:
                            content.seek(0)
                            width_px = int(block.get("width_in", 3) * 96)
                            align = block.get("alignment", "Left")
                            if align == "Center":
                                _, col, _ = st.columns([1, 2, 1])
                                col.image(content, width=width_px)
                            elif align == "Right":
                                _, col = st.columns([1, 1])
                                col.image(content, width=width_px)
                            else:
                                st.image(content, width=width_px)
                        except Exception:
                            st.write("🖼️ Image inserted")
                    else:
                        st.write("🖼️ Image inserted")
                elif block["type"] == "page_break":
                    st.markdown("---")
